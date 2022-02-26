from docx import Document
from docx.shared import Pt
document = Document()

Document.add_heading('', level='1')

p1 = document.add_paragraph('xxx')
p1.insert_paragraph('yyy')

Document.save('教案模板.docx')